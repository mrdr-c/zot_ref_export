# Standard library imports
from copy import deepcopy
import json
from time import sleep
from os import path

# 3rd-party package imports
from docx import Document
from pyzotero import zotero

# Global variables
EXIT_TIMER = 3


def main():
    # Loading config & setting locale
    config = get_config('config.json')
    locale = set_locale('locale.json', config['locale'])
    # pprint.pp(locale)

    # Saving messages and formatting dicts
    user_messages = locale['messages']
    formatting = locale['formatting']

    # Retrieving library
    zot = zotero.Zotero(config['library_id'],
                        config['library_type'],
                        config['api_key'])

    # Querying collection to export
    coll_name = (input(user_messages['prompt_coll_name']))  # TODO: make sure that UTF-8 encoding works fine
    coll_metadata = zot.collections(q=str.lower(coll_name))

    # checking if an unambiguous collection was found
    if len(coll_metadata) == 0:
        print(user_messages['err_no_coll_found'].format(coll_name=coll_name))
        sleep(EXIT_TIMER)
        exit(1)
    elif len(coll_metadata) > 1:
        print(user_messages['err_ambiguous_term'].format(coll_name=coll_name, len_coll_meta=len(coll_metadata)))
        sleep(EXIT_TIMER)
        exit(1)

    # Getting collection key - returns list of dicts. If one collection is returned, it is at position 0
    coll_key = coll_metadata[0]['key']

    # Checking for child collections
    if len(zot.collections_sub(coll_key)) > 0:
        print(user_messages['err_child_cols'].format(coll_name=coll_name))
        sleep(EXIT_TIMER)
        exit(1)

    # Getting all collection items as a list of dicts
    coll = zot.collection_items(coll_key)

    # TODO: Remove template entry if there is one

    # Getting the entries and saving them into a dict
    coll_entries = list()
    i = 0
    while i < len(coll):
        coll_entries.append(Entry(config, coll[i]))
        i += 1

    build_docx(coll_entries,
               coll_name,
               config['saving_location'],
               config['output_formatting'],
               formatting,
               user_messages)


# = = = = CLASSES = = = =

class Entry:
    """This is a class to hold the information of each entry as it is retrieved from the collection as an element.
    Pass a config dict and a raw entry dict to it.
    """

    def __init__(self, config, raw_entry):
        # Getting the fields to fetch
        fetch = dict()
        for field in config['fields_to_fetch']:
            fetch[field] = deepcopy(config['fields_to_fetch'][field])

        # saving content and the output order of elements in separate dicts
        self.data = dict()
        self.positions = dict()
        for element in fetch:
            self.data[element] = self.path_item(raw_entry, fetch[element]['path'])

            # Sorting the elements if their data entry isn't none
            if self.data[element]:
                self.positions[fetch[element]['position']] = element

        self.order = list()
        for key in self.positions:
            # Will skip positions that can't be interpreted
            try:
                self.order.append(int(key))
            except ValueError:
                pass

        # Sorting elements' positions (ascending order)
        self.order.sort()

        # Turning the info_string into a dict
        self.data['info_string'] = self.process_info_str(self.data['info_string'])

        # Turning the contractors dict list into a list
        contractor_list = list()
        for contractor in self.data['contractors']:
            contractor_list.append(contractor['name'])
        self.data['contractors'] = contractor_list

    def path_item(self, item, item_path):
        """Recurse to the bottom of a path (array) to retrieve a dict item"""
        # Exit condition
        if len(item_path) == 0:
            return item

        # Recursion logic
        item = item.get(item_path.pop(0))
        return self.path_item(item, item_path)

    @staticmethod
    def process_info_str(info_str, kv_delimiter='=', item_delimiter=';'):
        """Takes the info string and turns it into a dict"""
        # splitting string items
        items = info_str.replace('\n', '').split(item_delimiter)

        # splitting items into key/value pairs
        items_dict = dict()
        for item in items:
            k, v = item.split(kv_delimiter)
            v = v.strip()
            # Turning empty and 'none' strings into None type
            if v.lower() in ['none', '']:
                v = None
            items_dict[k.strip()] = v

        return items_dict


# = = = = FUNCTIONS = = = = 

def build_docx(entries_list, document_name, saving_location, formatting, locale_formatting, user_messages):
    """Builds document from entries list"""
    # TODO: refactor this into a modular, flexible behaviour defined within config.json
    #   - Define a title field as either None or a name to be found in locale.json
    #   - Make it deal with dicts and lists in a smart way
    #   - Handle singular and plural forms where applicable

    document = Document()

    for entry in entries_list:
        for ordinal in entry.order:
            field_name = entry.positions[ordinal]
            content = deepcopy(entry.data[field_name])

            # Catching None type
            if content:
                # Handling title
                if field_name == 'title':
                    document.add_heading(content, level=formatting['title_heading_lvl'])
                # Handling contractors
                elif field_name == 'contractors':
                    document.add_paragraph(f'{locale_formatting[field_name]}: {content[0]}')
                    if len(content) > 1:
                        for contractor in content[1:]:
                            document.add_run(f', {contractor}')
                # Handling info_strings:
                elif field_name == 'info_string':
                    for fn, ct in content.items():
                        document.add_paragraph(f'{fn}: {ct}')
                # Handling description
                elif field_name == 'description':
                    document.add_paragraph(content)
                # Handling other fields
                else:
                    document.add_paragraph(f'{locale_formatting[field_name]}: {content}')

    # Saving to downloads with specified name
    try:
        document.save(path.join(path.expanduser(saving_location), f'{document_name}.docx'))
    # Catching error caused by opened document
    except PermissionError:
        print(user_messages['err_opened_file'].format(filename=document_name))
        sleep(EXIT_TIMER)
        exit(1)

    print(user_messages['success'].format(filename=document_name, filepath=saving_location))
    sleep(EXIT_TIMER)
    return None


def get_config(filename):
    """Gets the config from the specified file and returns a dict of it. Handles missing file by creating template."""
    # TODO: Add additional checks to make sure that the fields 'library_id' and 'library_type' have the right format
    # TODO: Figure out how to store credentials securely

    try:
        with open(filename, mode='r', encoding='utf-8') as f:
            config = json.load(f)
    except FileNotFoundError:
        empty_config = {'library_id': None,
                        'library_type': None,
                        'api_key': None,
                        'locale': None,
                        'saving_location': '~/Downloads',
                        'fields_to_fetch': {
                            'title': {
                                'path': ['data', 'title'],
                                'position': 1
                            },
                            "contractors": {
                                'path': ['data', 'creators'],
                                'position': 2
                            },
                            'client': {
                                'path': ['data', 'publisher'],
                                'position': 3
                            },
                            'description': {
                                'path': ['data', 'abstractNote'],
                                'position': 5
                            },
                            'info_string': {
                                'path': ['data', 'extra'],
                                'position': 4
                            }
                        },
                        'output_formatting': {
                            'title_heading_lvl': 3
                        }
                        }

        with open(filename, 'w') as f:
            f.write(json.dumps(empty_config))
        print(f"'{filename}' not found in the execution directory. A blank file was created. Consult documentation.")
        sleep(EXIT_TIMER)
        exit(1)

    # assert isinstance(config, dict)
    return config


def set_locale(filename, loc, fallback_loc='en'):
    """Gets locale-specific messages and formatting from json file"""

    # Trying to open the specified file
    try:
        with open(filename, mode='r', encoding='utf-8') as f:
            locale = json.load(f)
    except FileNotFoundError:
        print(f"Missing '{filename}' file")
        sleep(EXIT_TIMER)
        exit(1)

    # Trying to retrieve the specified locale setting
    try:
        locale = locale[loc]
    except KeyError:
        # Trying to use the default [specified] fallback locale
        try:
            locale = locale[fallback_loc]
            print(f"Fallback locale '{fallback_loc}' set")
        except KeyError:
            print(f"fallback locale '{fallback_loc}' not found")
            sleep(EXIT_TIMER)
            exit(1)

    # assert isinstance(locale, dict)
    return locale


if __name__ == '__main__':
    main()
